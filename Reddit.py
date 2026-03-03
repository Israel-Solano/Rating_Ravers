#!/usr/bin/env python3
"""
reddit_best_comments.py

Fetch Reddit submissions, filter comments with score >= MIN_SCORE,
and export them (body, score, date) to a Word document.

The script accepts:
- A single Reddit URL
- A list of URLs from a text file (one URL per line) via --url-file
- A generic search query via --search, which is sent to a web search
  engine restricted to site:reddit.com, and the resulting Reddit
  URLs are processed.
"""

import argparse
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlsplit, urlunsplit, quote_plus

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm


# ----------------------------------------------------------------------
# Helper functions
# ----------------------------------------------------------------------

def _clean_url(url: str) -> str:
    """Remove query strings and fragments – Reddit's `.json` endpoint only needs the base URL."""
    parts = urlsplit(url)
    cleaned = parts._replace(query="", fragment="")
    return urlunsplit(cleaned)


def fetch_submission_json(url: str) -> dict:
    """Retrieve the JSON representation of a Reddit submission (/.json)."""
    url = _clean_url(url)

    if not url.endswith("/"):
        url += "/"
    json_url = url + ".json"
    headers = {"User-Agent": "reddit-best-comments/0.1 (by u/yourusername)"}
    # verify=False suppresses SSL warnings in environments without proper certs;
    # you may want to set it to True for production use.
    resp = requests.get(json_url, headers=headers, timeout=15)
    resp.raise_for_status()
    data = resp.json()
    if not isinstance(data, list) or len(data) < 2:
        raise ValueError("Unexpected JSON structure")
    return data[1]  # second element contains the comment tree


def walk_comments(comment_list, min_score, results):
    """Recursively walk the comment tree and collect comments with score >= min_score."""
    for comment in comment_list:
        if comment.get("kind") != "t1":
            continue
        data = comment["data"]
        body = data.get("body")
        if not body:
            continue

        score = data.get("score", 0)
        if score >= min_score:
            created_utc = data.get("created_utc")
            created_dt = datetime.fromtimestamp(created_utc, tz=timezone.utc)
            results.append(
                {
                    "author": data.get("author"),
                    "score": score,
                    "date": created_dt,
                    "body": body,
                }
            )

        replies = data.get("replies")
        if isinstance(replies, dict):
            children = replies.get("data", {}).get("children", [])
            walk_comments(children, min_score, results)


def _add_ai_instruction_page(doc: Document):
    """
    Add an AI instruction page at the very beginning of the document.
    This page tells an AI assistant how to summarize the Reddit comments.
    """
    # Title
    title = doc.add_heading("📋 AI Summarization Instructions", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document contains Reddit comments collected and ranked by upvote score. "
        "Please follow the instructions below to summarize this content."
    )

    doc.add_heading("Your Task", level=1)
    doc.add_paragraph(
        "Summarize the Reddit comments in this document into a clear, digestible report docx"
        "of the most recommended tips, advice, or insights shared by the community. Do not over-simplify."
        "If information is mentioned, it should be added unless it's noise. Do not worry about being too long."
    )

    doc.add_heading("Prioritization Rules", level=1)
    rules = [
        "WEIGHT BY UPVOTES: Comments with higher upvote scores (↑) represent stronger community consensus. "
        "Prioritize their content more heavily in your summary — a comment with 500 upvotes should carry "
        "significantly more weight than one with 10.",
        "IDENTIFY RECURRING THEMES: Group similar tips or pieces of advice together, even if they come "
        "from different comments. Note when multiple high-scoring comments agree on the same point.",
        "SURFACE THE BEST TIPS FIRST: Lead with the most upvoted, most agreed-upon recommendations. "
        "Place niche or lower-consensus advice toward the end.",
        "DISCARD NOISE: Ignore off-topic comments, jokes, or low-substance replies. Focus on actionable "
        "or informative content.",
        "PRESERVE SPECIFICITY: If a highly upvoted comment includes specific details (product names, "
        "numbers, steps, warnings), include those details in your summary — don't generalize them away.",
    ]
    for rule in rules:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rule)

    doc.add_heading("Output Format", level=1)
    doc.add_paragraph(
        "Produce a structured summary with the following sections:"
    )
    format_items = [
        "Key Takeaways — 3–5 bullet points capturing the highest-consensus tips.",
        "Top Recommendations — A ranked or grouped breakdown of the best advice, with brief explanations. "
        "Mention approximate upvote ranges where relevant (e.g. 'Several comments with 200+ upvotes suggest...').",
        "Additional Insights — Useful but less unanimous points worth noting.",
        "Warnings or Caveats — Any highly upvoted warnings, counterpoints, or 'don'ts' mentioned by the community.",
    ]
    for item in format_items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

    doc.add_heading("Tone & Style", level=1)
    doc.add_paragraph(
        "Write in clear, plain English. Be concise. Avoid copying comments verbatim — paraphrase and synthesize. "
        "The goal is to save the reader time by distilling community wisdom into actionable insights."
    )

    # Separator note
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("─" * 40 + "  BEGIN REDDIT COMMENTS BELOW  " + "─" * 40)
    run.bold = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def _add_comments_to_doc(doc: Document, comments: list, url: str):
    """Append a block of comments for a single Reddit URL."""
    doc.add_heading(f"Comments from {url}", level=1)
    doc.add_paragraph(
        f"Fetched on {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M:%SZ')}"
    )
    doc.add_page_break()

    for i, cm in enumerate(comments, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        header = (
            f"u/{cm['author']} ‑ {cm['score']} ↑ ‑ "
            f"{cm['date'].strftime('%Y-%m-%d %H:%M UTC')}"
        )
        p.add_run(header).italic = True
        p.add_run("\n")
        body_run = p.add_run(cm["body"])
        body_run.font.size = Pt(10)
        p.add_run("\n\n")  # spacing between comments


def build_or_append_docx(comments: list, output_path: Path, url: str):
    """Load an existing doc (or create a new one) and append the supplied comments."""
    if output_path.exists():
        doc = Document(output_path)
    else:
        doc = Document()
        # Add AI instruction page first
        _add_ai_instruction_page(doc)

        doc.add_heading("Reddit Best Comments", level=0)
        doc.add_paragraph(
            f"Generated on {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M:%SZ')}"
        )
        doc.add_page_break()

    _add_comments_to_doc(doc, comments, url)
    doc.save(output_path)


def _process_one_url(url: str, min_score: int, output_path: Path):
    """Fetch, filter, and write comments for a single URL."""
    cleaned_url = _clean_url(url)
    print(f"\nFetching JSON for {cleaned_url} …")
    comment_json = fetch_submission_json(url)  # internal clean again

    top_level = comment_json.get("data", {}).get("children", [])
    filtered = []
    # print("Walking comment tree …")
    walk_comments(top_level, min_score, filtered)

    if not filtered:
        print(f"No comments with score >= {min_score} found for {cleaned_url}.")
        return

    filtered.sort(key=lambda x: x["score"], reverse=True)
    print(
        f"Appending {len(filtered)} comment(s) from {cleaned_url} to {output_path} …"
    )
    build_or_append_docx(filtered, output_path, cleaned_url)


# ----------------------------------------------------------------------
# New helper: use a web search engine restricted to Reddit
# ----------------------------------------------------------------------
def search_reddit_via_web(query: str, limit: int = 25) -> list[str]:
    """Use Google search restricted to Reddit comment threads."""
    full_query = f'"{query}" site:reddit.com/r/ "comments" '  # Precise pattern
    encoded_q = quote_plus(full_query)
    
    # Google search (mobile endpoint is more scraper-friendly)
    url = f"https://www.bing.com/search?q={encoded_q}&count={limit}"

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    
    print(f"Searching Google: {full_query}")
    resp = requests.get(url, headers=headers, timeout=15)
    resp.raise_for_status()
    
    # Simple regex to grab Reddit URLs from Google's result links
    import re
    reddit_urls = re.findall(r'https://www\.reddit\.com/r/[^\s"]+/comments/[^\s"]+', resp.text)
    
    urls = []
    for url in reddit_urls[:limit]:
        cleaned = _clean_url(url)
        if cleaned not in urls:
            urls.append(cleaned)
            print(f"✓ Found: {cleaned}")
    
    print(f"Final count: {len(urls)} URLs")
    return urls


# ----------------------------------------------------------------------
# Main entry point
# ----------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Export Reddit comments with high scores."
    )
    parser.add_argument(
        "url",
        nargs="?",
        help="URL of a Reddit submission (e.g. https://www.reddit.com/r/.../comments/xxxx/). "
             "If omitted, you must provide --url-file or --search.",
    )
    parser.add_argument(
        "--url-file",
        type=Path,
        help="Path to a text file containing one Reddit URL per line.",
    )
    parser.add_argument(
        "-s", "--search",
        metavar="QUERY",
        help="Search the web for Reddit submissions matching QUERY and process the results.",
    )
    parser.add_argument(
        "--search-limit",
        type=int,
        default=25,
        help="Maximum number of search results to process (default 25).",
    )
    parser.add_argument(
        "--min-score",
        type=int,
        default=3,
        help="Minimum up‑vote count to keep a comment (default: 5).",
    )
    parser.add_argument(
        "--output",
        default="reddit_best_comments.docx",
        help="Output Word document filename (default: reddit_best_comments.docx).",
    )
    args = parser.parse_args()

    # --------------------------------------------------------------
    # Validate mutually exclusive inputs
    # --------------------------------------------------------------
    input_modes = sum(bool(x) for x in (args.url, args.url_file, args.search))
    if input_modes == 0:
        parser.error("Provide either a URL, --url-file, or --search.")
    if input_modes > 1:
        parser.error("Only one of URL, --url-file, or --search may be supplied.")

    # --------------------------------------------------------------
    # Build the list of URLs to process
    # --------------------------------------------------------------
    url_list: list[str] = []

    if args.url_file:
        try:
            with args.url_file.open("r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    url_list.append(line)
        except OSError as e:
            print(f"Error reading URL file: {e}", file=sys.stderr)
            sys.exit(1)

    elif args.search:
        print(f"🔎 Searching the web for Reddit threads about '{args.search}' …")
        try:
            found = search_reddit_via_web(
                query=args.search,
                limit=args.search_limit,
            )
        except Exception as e:
            print(f"Search failed: {e}", file=sys.stderr)
            sys.exit(1)

        if not found:
            print("No Reddit submissions matched the query.", file=sys.stderr)
            sys.exit(0)

        print(f"✅ Found {len(found)} Reddit submissions; will process them.")
        url_list.extend(found)

    else:  # single positional URL
        url_list.append(args.url)

    # --------------------------------------------------------------
    # Process each URL (continue on error so one bad URL doesn't stop the batch)
    # --------------------------------------------------------------
    output_path = Path(args.output)
    for url in url_list:
        try:
            _process_one_url(url, args.min_score, output_path)
        except KeyboardInterrupt:
            print("\nInterrupted by user.", file=sys.stderr)
            sys.exit(1)
        except Exception as e:
            print(f"Error processing {url}: {e}", file=sys.stderr)
            continue

    print("\nAll done!")


if __name__ == "__main__":
    main()